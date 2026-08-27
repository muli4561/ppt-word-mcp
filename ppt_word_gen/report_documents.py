# coding=utf-8
"""DOCX 输入预检、证据抽取、模板渲染与确定性质量检查。"""
import hashlib
import json
import posixpath
import re
import shutil
import subprocess
import tempfile
import zipfile
from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docxtpl import DocxTemplate

from .config import MAX_REPORT_SOURCE_CHARS
from .pptmaster import convert_source
from .report_models import REPORT_TYPES, ReportBlock, ReportCreate, ReportSpec, WordFormatOptions


_SAFE_NAME = re.compile(r"[^0-9A-Za-z_.\-]")
_TOKEN_RE = re.compile(r"{{\s*(?:[ptrc]\s+)?([A-Za-z_][A-Za-z0-9_]*)\s*}}")
_NUMBER_RE = re.compile(r"(?<![\w.])[+-]?\d+(?:\.\d+)?%?")
_REL_NS = {"r": "http://schemas.openxmlformats.org/package/2006/relationships"}
_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".emf", ".wmf"}
STRICT_WORD_FORMAT_PROFILE_ID = "cid629-joint-simulation-v1.5"


def safe_filename(filename: str, fallback: str) -> str:
    name = _SAFE_NAME.sub("_", Path(filename).name).strip("._")
    suffix = Path(filename).suffix.lower()
    return (name or f"{fallback}{suffix}")[:180]


def write_upload_atomic(directory: Path, data: bytes, filename: str, fallback: str) -> Path:
    directory.mkdir(parents=True, exist_ok=True)
    path = directory / safe_filename(filename, fallback)
    temp_path = path.with_suffix(path.suffix + ".part")
    temp_path.write_bytes(data)
    if temp_path.stat().st_size != len(data):
        temp_path.unlink(missing_ok=True)
        raise ValueError("上传文件落盘大小不一致")
    temp_path.replace(path)
    return path


def _is_symlink(info: zipfile.ZipInfo) -> bool:
    return ((info.external_attr >> 16) & 0o170000) == 0o120000


def _relationship_source_dir(rels_name: str) -> str:
    if rels_name == "_rels/.rels":
        return ""
    parent = posixpath.dirname(posixpath.dirname(rels_name))
    return parent


def validate_docx_package(path: Path) -> Dict:
    path = Path(path)
    if path.suffix.lower() not in {".docx", ".dotx"}:
        raise ValueError("参考模板必须是 .docx 或 .dotx 文件")
    if not zipfile.is_zipfile(path):
        raise ValueError("文件不是有效的 DOCX ZIP 包")

    errors: List[str] = []
    warnings: List[str] = []
    with zipfile.ZipFile(path) as archive:
        infos = archive.infolist()
        names = set(archive.namelist())
        required = {"[Content_Types].xml", "_rels/.rels", "word/document.xml"}
        missing = sorted(required - names)
        if missing:
            errors.append("缺少必要部件: " + ", ".join(missing))
        if len(infos) > 5000:
            errors.append("DOCX 内部文件数量异常")
        total_uncompressed = sum(info.file_size for info in infos)
        if total_uncompressed > 250 * 1024 * 1024:
            errors.append("DOCX 解压后体积超过 250MB")
        if any(_is_symlink(info) for info in infos):
            errors.append("DOCX 包含不允许的符号链接条目")

        for rels_name in (name for name in names if name.endswith(".rels")):
            try:
                root = ET.fromstring(archive.read(rels_name))
            except ET.ParseError as exc:
                errors.append(f"关系文件 XML 损坏: {rels_name}: {exc}")
                continue
            source_dir = _relationship_source_dir(rels_name)
            for rel in root.findall("r:Relationship", _REL_NS):
                if rel.get("TargetMode") == "External":
                    continue
                target = rel.get("Target", "")
                if not target:
                    continue
                resolved = posixpath.normpath(posixpath.join(source_dir, target.lstrip("/")))
                if resolved.startswith("../") or resolved not in names:
                    errors.append(f"关系目标不存在: {rels_name} -> {target}")

        media = [name for name in names if name.startswith("word/media/") and not name.endswith("/")]
        for name in media:
            suffix = Path(name).suffix.lower()
            if suffix not in _IMAGE_SUFFIXES:
                warnings.append(f"未知媒体类型: {name}")
                continue
            data = archive.read(name)[:16]
            known = (
                data.startswith(b"\x89PNG\r\n\x1a\n")
                or data.startswith(b"\xff\xd8\xff")
                or data.startswith((b"GIF87a", b"GIF89a", b"BM", b"II*\x00", b"MM\x00*"))
                or suffix in {".emf", ".wmf"}
            )
            if not known:
                errors.append(f"媒体文件签名异常: {name}")

    if errors:
        raise ValueError("；".join(errors[:12]))
    return {
        "ok": True,
        "entries": len(infos),
        "uncompressed_bytes": total_uncompressed,
        "media_count": len(media),
        "warnings": warnings,
    }


def _cjk_count(text: str) -> int:
    return sum(1 for char in text if "\u3400" <= char <= "\u9fff")


def normalize_mojibake(text: str) -> Tuple[str, bool]:
    """修复常见的 GBK 字节被错误解释为 Latin-1 的文本。"""
    changed = False
    output: List[str] = []
    for line in text.splitlines(keepends=True):
        try:
            repaired = line.encode("latin1").decode("gb18030")
        except (UnicodeEncodeError, UnicodeDecodeError):
            output.append(line)
            continue
        if _cjk_count(repaired) >= _cjk_count(line) + 2:
            output.append(repaired)
            changed = True
        else:
            output.append(line)
    return "".join(output), changed


def extract_source(source_path: Optional[Path], instructions: str) -> Tuple[str, List[Path], Dict]:
    parts: List[str] = []
    images: List[Path] = []
    metadata: Dict = {"source_file": None, "encoding_repaired": False}
    if instructions.strip():
        parts.append("# 用户要求\n\n" + instructions.strip())

    if source_path is not None:
        metadata["source_file"] = source_path.name
        metadata["sha256"] = hashlib.sha256(source_path.read_bytes()).hexdigest()
        suffix = source_path.suffix.lower()
        if suffix == ".docx":
            metadata["docx_preflight"] = validate_docx_package(source_path)

        if suffix in {".md", ".txt"}:
            converted = source_path.read_text(encoding="utf-8", errors="replace")
        else:
            convert_source(str(source_path), cwd=source_path.parent)
            exact = source_path.with_suffix(".md")
            candidates = [exact] if exact.is_file() else sorted(
                source_path.parent.glob(f"{source_path.stem}*.md"),
                key=lambda item: item.stat().st_mtime,
                reverse=True,
            )
            if not candidates:
                raise RuntimeError("源资料转换完成，但未找到 Markdown 输出")
            converted = candidates[0].read_text(encoding="utf-8", errors="replace")

        converted, repaired = normalize_mojibake(converted)
        metadata["encoding_repaired"] = repaired
        parts.append(f"# 来源文件：{source_path.name}\n\n{converted}")

        media_dir = source_path.parent / f"{source_path.stem}_files"
        if media_dir.is_dir():
            images = [
                item for item in sorted(media_dir.rglob("*"))
                if item.is_file() and item.suffix.lower() in _IMAGE_SUFFIXES
            ]

    combined = "\n\n".join(parts).strip()
    truncated = len(combined) > MAX_REPORT_SOURCE_CHARS
    metadata["source_chars"] = len(combined)
    metadata["source_truncated"] = truncated
    metadata["images"] = [item.name for item in images]
    return combined[:MAX_REPORT_SOURCE_CHARS], images, metadata


def build_evidence_manifest(source_text: str, source_metadata: Dict) -> Dict:
    entries = []
    for raw in source_text.splitlines():
        text = re.sub(r"\s+", " ", raw).strip(" #\t")
        if len(text) < 8:
            continue
        entries.append({
            "id": f"E{len(entries) + 1:04d}",
            "text": text[:1600],
            "numbers": _NUMBER_RE.findall(text),
        })
        if len(entries) >= 800:
            break
    return {
        "source": source_metadata,
        "entries": entries,
        "evidence_count": len(entries),
    }


def _set_run_font(run, name: str = "Microsoft YaHei", size: Optional[float] = None) -> None:
    run.font.name = name
    if size:
        run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def _configure_default_styles(document: Document) -> None:
    styles = document.styles
    for style_name, size, bold in (
        ("Normal", 10.5, False),
        ("Title", 24, True),
        ("Subtitle", 12, False),
        ("Heading 1", 16, True),
        ("Heading 2", 14, True),
        ("Heading 3", 12, True),
    ):
        try:
            style = styles[style_name]
        except KeyError:
            continue
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.bold = bold
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _set_style_font(style, font_name: str, size_pt: float) -> None:
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), font_name)


def _ensure_multilevel_numbering(document: Document, numbering_style: str) -> None:
    root = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId"), "0"))
        for node in root.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"), "0")) for node in root.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    number_format = "chineseCounting" if numbering_style == "chinese" else "decimal"
    style_names = ["Heading1", "Heading2", "Heading3"]
    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), number_format)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), ".".join(f"%{index}" for index in range(1, level + 2)) + ".")
        paragraph_style = OxmlElement("w:pStyle")
        paragraph_style.set(qn("w:val"), style_names[level])
        lvl.append(start)
        lvl.append(num_fmt)
        lvl.append(text)
        lvl.append(paragraph_style)
        abstract.append(lvl)
    root.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    number.append(abstract_ref)
    root.append(number)

    for level, name in enumerate(("Heading 1", "Heading 2", "Heading 3")):
        try:
            style = document.styles[name]
        except KeyError:
            continue
        ppr = style._element.get_or_add_pPr()
        existing = ppr.find(qn("w:numPr"))
        if existing is not None:
            ppr.remove(existing)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level))
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_node)
        ppr.insert(0, num_pr)


def _apply_word_format(document: Document, options: WordFormatOptions) -> None:
    styles = document.styles
    normal = styles["Normal"]
    _set_style_font(normal, options.body_font, options.body_size_pt)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = options.line_spacing
    normal.paragraph_format.first_line_indent = Cm(options.first_line_indent_chars * 0.3704166667)
    for name, font, size in (
        ("Heading 1", options.heading1_font, options.heading1_size_pt),
        ("Heading 2", options.heading2_font, options.heading2_size_pt),
        ("Heading 3", options.heading3_font, options.heading3_size_pt),
    ):
        try:
            _set_style_font(styles[name], font, size)
        except KeyError:
            continue
    _ensure_multilevel_numbering(document, options.numbering_style)


def _clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _add_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "打开 Word 后更新目录"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, placeholder, end):
        run._r.append(element)
    settings = document.settings._element
    if settings.find(qn("w:updateFields")) is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)


def _add_metadata_table(document: Document, spec: ReportSpec, set_default_font: bool = True) -> None:
    rows = [
        ("报告类型", REPORT_TYPES.get(spec.report_type, spec.report_type)),
        ("项目名称", spec.project_name or "未提供"),
        ("文档版本", spec.document_version or "v1.0"),
        ("编制人", spec.author or "未提供"),
        ("日期", spec.report_date or date.today().isoformat()),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for index, (label, value) in enumerate(rows):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = value
        for run in table.cell(index, 0).paragraphs[0].runs:
            run.bold = True
            if set_default_font:
                _set_run_font(run)
        for run in table.cell(index, 1).paragraphs[0].runs:
            if set_default_font:
                _set_run_font(run)


def _add_table(document: Document, block: ReportBlock, set_default_font: bool = True) -> None:
    width = len(block.headers)
    table = document.add_table(rows=1, cols=width)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for index, header in enumerate(block.headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            if set_default_font:
                _set_run_font(run)
    for raw_row in block.rows:
        cells = table.add_row().cells
        for index in range(width):
            cells[index].text = str(raw_row[index]) if index < len(raw_row) else ""
            for run in cells[index].paragraphs[0].runs:
                if set_default_font:
                    _set_run_font(run)


def _add_with_style(document: Document, text: str, style: str, fallback_prefix: str = ""):
    try:
        document.styles[style]
    except KeyError:
        return document.add_paragraph(fallback_prefix + text)
    return document.add_paragraph(text, style=style)


def _resolve_image(image_root: Path, image_name: str) -> Path:
    root = image_root.resolve()
    candidates = [item for item in root.rglob(Path(image_name).name) if item.is_file()]
    if not candidates:
        raise ValueError(f"报告引用的图片不存在: {image_name}")
    resolved = candidates[0].resolve()
    if root not in resolved.parents:
        raise ValueError("图片路径超出任务目录")
    return resolved


def _append_spec(
    document: Document,
    spec: ReportSpec,
    image_root: Path,
    include_title: bool,
    apply_default_formatting: bool = True,
) -> None:
    if apply_default_formatting:
        _configure_default_styles(document)
    if include_title:
        title = document.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(spec.title)
        if apply_default_formatting:
            _set_run_font(title_run, size=24)
        if spec.subtitle:
            subtitle = document.add_paragraph(style="Subtitle")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.add_run(spec.subtitle)
            if apply_default_formatting:
                _set_run_font(subtitle_run, size=12)
        document.add_paragraph()
        _add_metadata_table(document, spec, apply_default_formatting)
        document.add_page_break()
        document.add_heading("目录", level=1)
        _add_toc(document)
        document.add_page_break()

    document.add_heading("摘要", level=1)
    document.add_paragraph(spec.executive_summary)
    for section in spec.sections:
        document.add_heading(section.heading, level=section.level)
        for block in section.blocks:
            if block.type == "paragraph":
                document.add_paragraph(block.text)
            elif block.type == "bullets":
                for item in block.items:
                    _add_with_style(document, item, "List Bullet", "• ")
            elif block.type == "table":
                _add_table(document, block, apply_default_formatting)
                if block.caption:
                    _add_with_style(document, block.caption, "Caption")
            elif block.type == "image":
                image = _resolve_image(image_root, block.image_name)
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(str(image), width=Cm(15.5))
                if block.caption:
                    caption = _add_with_style(document, block.caption, "Caption")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif block.type == "page_break":
                document.add_page_break()

    if spec.conclusions:
        document.add_heading("结论", level=1)
        for item in spec.conclusions:
            _add_with_style(document, item, "List Bullet", "• ")
    if spec.risks:
        document.add_heading("风险与待确认项", level=1)
        for item in spec.risks:
            _add_with_style(document, item, "List Bullet", "• ")


def _set_cell_text(table, row: int, column: int, value: str) -> None:
    if row >= len(table.rows) or column >= len(table.rows[row].cells):
        return
    cell = table.cell(row, column)
    cell.text = value


def _append_body_element(document: Document, element) -> None:
    body = document._element.body
    section_properties = body.find(qn("w:sectPr"))
    if section_properties is None:
        body.append(element)
    else:
        section_properties.addprevious(element)


def _render_strict_reference(document: Document, spec: ReportSpec, image_root: Path) -> None:
    """保留 CID629 的封面/更新记录和样式，清除示例正文后填充报告。"""
    if len(document.tables) < 2:
        raise ValueError("固定 Word 模板缺少封面表格或文件更新记录表格")
    cover = document.tables[0]
    _set_cell_text(cover, 0, 4, spec.project_name or REPORT_TYPES.get(spec.report_type, spec.report_type))
    _set_cell_text(cover, 1, 4, REPORT_TYPES.get(spec.report_type, spec.report_type))
    _set_cell_text(cover, 2, 4, "")
    _set_cell_text(cover, 3, 4, spec.document_version or "v1.0")
    _set_cell_text(cover, 4, 0, spec.title)
    _set_cell_text(cover, 5, 4, spec.author or "未提供")
    for row in (6, 7, 8):
        _set_cell_text(cover, row, 4, "")

    history = document.tables[1]
    while len(history.rows) > 2:
        history._tbl.remove(history.rows[-1]._tr)
    if len(history.rows) > 1:
        values = [spec.report_date or date.today().isoformat(), spec.author or "未提供", spec.document_version, "AI 生成初版"]
        for index, value in enumerate(values):
            _set_cell_text(history, 1, index, value)

    cover_xml = deepcopy(cover._tbl)
    history_xml = deepcopy(history._tbl)
    history_title = None
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "文件更新记录":
            history_title = deepcopy(paragraph._p)
            break

    _clear_body(document)
    _append_body_element(document, cover_xml)
    document.add_page_break()
    if history_title is not None:
        _append_body_element(document, history_title)
    else:
        document.add_paragraph("文件更新记录")
    _append_body_element(document, history_xml)
    document.add_page_break()
    document.add_heading("目录", level=1)
    _add_toc(document)
    document.add_page_break()
    _append_spec(
        document,
        spec,
        image_root,
        include_title=False,
        apply_default_formatting=False,
    )


def _template_tokens(path: Path) -> List[str]:
    tokens = set()
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            text = archive.read(name).decode("utf-8", errors="ignore")
            tokens.update(_TOKEN_RE.findall(text))
    return sorted(tokens)


def render_report(
    spec: ReportSpec,
    output_path: Path,
    image_root: Path,
    reference_template: Optional[Path] = None,
    strict_reference: bool = False,
    format_options: Optional[WordFormatOptions] = None,
    format_profile_id: str = "",
) -> Dict:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    mode = "default"
    tokens: List[str] = []
    options = format_options or WordFormatOptions()

    if reference_template is None:
        document = Document()
        section = document.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.3)
        _configure_default_styles(document)
        _apply_word_format(document, options)
        _append_spec(document, spec, image_root, include_title=True, apply_default_formatting=False)
        document.save(output_path)
    else:
        validate_docx_package(reference_template)
        tokens = _template_tokens(reference_template)
        if strict_reference:
            mode = "strict-cid629-reference"
            document = Document(reference_template)
            _apply_word_format(document, options)
            _render_strict_reference(document, spec, image_root)
            document.save(output_path)
        elif "report_body" in tokens:
            mode = "prepared-template"
            template = DocxTemplate(str(reference_template))
            subdocument = template.new_subdoc()
            _apply_word_format(subdocument, options)
            _append_spec(
                subdocument,
                spec,
                image_root,
                include_title=False,
                apply_default_formatting=False,
            )
            context = spec.model_dump(mode="json")
            context.update({
                "report_type_name": REPORT_TYPES.get(spec.report_type, spec.report_type),
                "report_body": subdocument,
            })
            template.render(context, autoescape=True)
            template.save(output_path)
        elif tokens:
            mode = "token-template"
            template = DocxTemplate(str(reference_template))
            context = spec.model_dump(mode="json")
            context["report_type_name"] = REPORT_TYPES.get(spec.report_type, spec.report_type)
            template.render(context, autoescape=True)
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
                temp_path = Path(handle.name)
            try:
                template.save(temp_path)
                document = Document(temp_path)
                _apply_word_format(document, options)
                _append_spec(
                    document,
                    spec,
                    image_root,
                    include_title="title" not in tokens,
                    apply_default_formatting=False,
                )
                document.save(output_path)
            finally:
                temp_path.unlink(missing_ok=True)
        else:
            mode = "style-reference"
            document = Document(reference_template)
            _clear_body(document)
            _apply_word_format(document, options)
            _append_spec(
                document,
                spec,
                image_root,
                include_title=True,
                apply_default_formatting=False,
            )
            document.save(output_path)

    metadata = {
        "mode": mode,
        "template_tokens": tokens,
        "format_profile_id": format_profile_id or (
            STRICT_WORD_FORMAT_PROFILE_ID if strict_reference else "custom-template"
        ),
        "format_options": options.model_dump(mode="json"),
    }
    return metadata


def _all_document_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    for section in document.sections:
        values.extend(paragraph.text for paragraph in section.header.paragraphs)
        values.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(values)


def _iter_blocks(spec: ReportSpec) -> Iterable[ReportBlock]:
    for section in spec.sections:
        yield from section.blocks


def _effective_east_asia_font(style) -> str:
    current = style
    visited = set()
    while current is not None and current.style_id not in visited:
        visited.add(current.style_id)
        rpr = current._element.rPr
        fonts = rpr.rFonts if rpr is not None else None
        if fonts is not None:
            value = fonts.get(qn("w:eastAsia"))
            if value:
                return value
        current = current.base_style
    return ""


def _strict_format_checks(
    document: Document,
    docx_path: Path,
    options: WordFormatOptions,
) -> Tuple[Dict, List[str]]:
    errors: List[str] = []
    styles = document.styles
    normal = styles["Normal"]
    heading1 = styles["Heading 1"]
    heading2 = styles["Heading 2"]

    checks = {
        "profile_id": STRICT_WORD_FORMAT_PROFILE_ID,
        "normal_font": _effective_east_asia_font(normal),
        "normal_size_pt": normal.font.size.pt if normal.font.size else None,
        "normal_justified": normal.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY,
        "normal_line_spacing": normal.paragraph_format.line_spacing,
        "normal_first_line_indent_cm": (
            normal.paragraph_format.first_line_indent.cm
            if normal.paragraph_format.first_line_indent is not None
            else None
        ),
        "heading1_font": _effective_east_asia_font(heading1),
        "heading1_size_pt": heading1.font.size.pt if heading1.font.size else None,
        "heading2_font": _effective_east_asia_font(heading2),
        "heading2_size_pt": heading2.font.size.pt if heading2.font.size else None,
        "heading1_paragraphs": sum(p.style.name == "Heading 1" for p in document.paragraphs),
        "heading2_paragraphs": sum(p.style.name == "Heading 2" for p in document.paragraphs),
    }
    heading1_num = heading1._element.pPr.numPr if heading1._element.pPr is not None else None
    heading2_num = heading2._element.pPr.numPr if heading2._element.pPr is not None else None
    checks["heading1_numbered"] = heading1_num is not None
    checks["heading2_numbered"] = heading2_num is not None

    with zipfile.ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        settings_xml = archive.read("word/settings.xml").decode("utf-8", errors="ignore")
    checks["toc_levels_1_to_3"] = 'TOC \\o &quot;1-3&quot;' in document_xml or 'TOC \\o "1-3"' in document_xml
    checks["update_fields_on_open"] = "updateFields" in settings_xml

    expected = [
        (checks["normal_font"] == options.body_font, f"正文样式不是 {options.body_font}"),
        (checks["normal_size_pt"] == options.body_size_pt, f"正文字号不是 {options.body_size_pt:g}pt"),
        (checks["normal_justified"], "正文不是两端对齐"),
        (checks["normal_line_spacing"] == options.line_spacing, f"正文不是 {options.line_spacing:g} 倍行距"),
        (
            checks["normal_first_line_indent_cm"] is not None
            and abs(
                checks["normal_first_line_indent_cm"]
                - options.first_line_indent_chars * 0.3704166667
            ) <= 0.03,
            "正文首行缩进不符合确认格式",
        ),
        (checks["heading1_font"] == options.heading1_font, f"一级标题不是 {options.heading1_font}"),
        (checks["heading1_size_pt"] == options.heading1_size_pt, f"一级标题字号不是 {options.heading1_size_pt:g}pt"),
        (checks["heading2_font"] == options.heading2_font, f"二级标题不是 {options.heading2_font}"),
        (checks["heading2_size_pt"] == options.heading2_size_pt, f"二级标题字号不是 {options.heading2_size_pt:g}pt"),
        (checks["heading1_paragraphs"] >= 1, "报告缺少一级标题"),
        (checks["heading2_paragraphs"] >= 1, "报告缺少二级标题"),
        (checks["heading1_numbered"], "一级标题未绑定多级编号"),
        (checks["heading2_numbered"], "二级标题未绑定多级编号"),
        (checks["toc_levels_1_to_3"], "报告缺少 1–3 级自动目录"),
        (checks["update_fields_on_open"], "报告未设置打开时更新目录"),
    ]
    errors.extend(message for passed, message in expected if not passed)
    return checks, errors


def _render_pdf_if_available(docx_path: Path, output_dir: Path) -> Dict:
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        return {"status": "skipped", "reason": "LibreOffice 未安装"}
    proc = subprocess.run(
        [executable, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=180,
    )
    pdf = output_dir / f"{docx_path.stem}.pdf"
    if proc.returncode != 0 or not pdf.is_file():
        return {"status": "failed", "error": (proc.stderr or proc.stdout)[-1000:]}
    return {"status": "rendered", "pdf": str(pdf), "visual_inspection": "not_automated"}


def validate_rendered_report(
    docx_path: Path,
    spec: ReportSpec,
    evidence: Dict,
    render_metadata: Dict,
) -> Dict:
    package = validate_docx_package(docx_path)
    document = Document(docx_path)
    text = _all_document_text(document)
    errors: List[str] = []
    warnings: List[str] = list(package.get("warnings", []))

    with zipfile.ZipFile(docx_path) as archive:
        unresolved = set()
        for name in archive.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                unresolved.update(_TOKEN_RE.findall(archive.read(name).decode("utf-8", errors="ignore")))
    if unresolved:
        errors.append("存在未替换模板变量: " + ", ".join(sorted(unresolved)))
    if spec.title not in text:
        errors.append("输出文档中未找到报告标题")
    if len(text.strip()) < 200:
        errors.append("输出文档正文内容过少")

    strict_format = None
    if (
        render_metadata.get("format_options")
        and render_metadata.get("mode") not in {"prepared-template", "token-template"}
    ):
        expected_format = WordFormatOptions.model_validate(render_metadata["format_options"])
        strict_format, strict_errors = _strict_format_checks(document, docx_path, expected_format)
        errors.extend(strict_errors)

    evidence_map = {entry["id"]: entry for entry in evidence.get("entries", [])}
    for block in _iter_blocks(spec):
        missing_ids = sorted(set(block.evidence_ids) - set(evidence_map))
        if missing_ids:
            errors.append("引用了不存在的证据: " + ", ".join(missing_ids))
        if evidence_map and block.type != "page_break" and not block.evidence_ids:
            errors.append(f"内容块缺少证据引用: {block.type}")
        if not block.evidence_ids:
            continue
        block_text = " ".join([
            block.text,
            *block.items,
            *block.headers,
            *(str(cell) for row in block.rows for cell in row),
        ])
        referenced = " ".join(evidence_map[item]["text"] for item in block.evidence_ids if item in evidence_map)
        unmatched = [number for number in _NUMBER_RE.findall(block_text) if number not in referenced]
        if unmatched:
            warnings.append(f"数值未在引用证据中逐字匹配: {', '.join(sorted(set(unmatched)))}")

    pdf_render = _render_pdf_if_available(docx_path, docx_path.parent)
    if pdf_render.get("status") == "failed":
        warnings.append("LibreOffice PDF 渲染失败")

    report = {
        "status": "failed" if errors else ("passed-with-warnings" if warnings else "passed"),
        "errors": errors,
        "warnings": warnings,
        "package": package,
        "structure": {
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "sections": len(document.sections),
            "characters": len(text),
        },
        "template": render_metadata,
        "strict_format": strict_format,
        "pdf_render": pdf_render,
    }
    if errors:
        raise ValueError("Word 报告质量检查失败: " + "；".join(errors[:10]))
    return report


def write_json(path: Path, value: Dict) -> None:
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
