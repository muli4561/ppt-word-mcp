import tempfile
import time
import unittest
from pathlib import Path

from docx import Document

from ppt_word_gen.report_models import ReportCreate
from ppt_word_gen.report_tasks import (
    ReportIdempotencyConflict,
    ReportTaskManager,
    report_manager as application_report_manager,
)
from ppt_word_gen.task_store import TaskStore


class ReportTaskManagerTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory(prefix="report-task-")
        root = Path(self.temp_dir.name)
        self.root = root
        self.manager = ReportTaskManager(
            max_workers=1,
            max_queued=2,
            store=TaskStore(root / "tasks.db"),
            output_dir=root / "output",
            mock_llm=True,
        )

    def tearDown(self):
        self.manager.shutdown(wait=True)
        self.temp_dir.cleanup()

    def _wait(self, task_id: str, timeout: float = 8):
        deadline = time.time() + timeout
        while time.time() < deadline:
            info = self.manager.get(task_id)
            if info.status in {"success", "failed", "cancelled", "interrupted"}:
                return info
            time.sleep(0.05)
        self.fail("report task did not finish")

    def test_mock_task_generates_docx_without_source_or_template(self):
        request = ReportCreate(
            report_type="delivery",
            instructions="生成 Agent 交付说明，区分事实和待确认项。",
            project_name="示例 Agent",
        )
        task_id, reused = self.manager.submit(request, idempotency_key="report-1")
        info = self._wait(task_id)

        self.assertFalse(reused)
        self.assertEqual("success", info.status, info.error)
        result = Path(self.manager.get_result_path(task_id))
        self.assertTrue(result.is_file())
        self.assertEqual(".docx", result.suffix)
        self.assertTrue((result.parent / "report_spec.json").is_file())
        self.assertTrue((result.parent / "validation.json").is_file())

    def test_template_bytes_are_part_of_idempotency_hash(self):
        request = ReportCreate(instructions="生成测试报告")
        path = self.root / "template.docx"
        doc = Document()
        doc.add_paragraph("模板 A")
        doc.save(path)
        template_a = (path.read_bytes(), "template.docx")
        doc.add_paragraph("模板 B 的新增内容")
        doc.save(path)
        template_b = (path.read_bytes(), "template.docx")
        task_id, _ = self.manager.submit(request, template_upload=template_a, idempotency_key="same")
        with self.assertRaises(ReportIdempotencyConflict):
            self.manager.submit(request, template_upload=template_b, idempotency_key="same")
        self._wait(task_id)


def tearDownModule():
    application_report_manager.shutdown()


if __name__ == "__main__":
    unittest.main()
