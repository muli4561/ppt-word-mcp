import tempfile
import unittest
from pathlib import Path

from docx import Document

from ppt_word_gen.config import WORD_REPORT_TEMPLATE_PATH
from ppt_word_gen.report_agent import build_mock_spec
from ppt_word_gen.report_documents import (
    _add_with_style,
    build_evidence_manifest,
    normalize_mojibake,
    render_report,
    validate_docx_package,
    validate_rendered_report,
)
from ppt_word_gen.report_models import ReportCreate, WordFormatOptions


class ReportDocumentTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory(prefix="report-docx-")
        self.root = Path(self.temp_dir.name)
        self.request = ReportCreate(
            title="联合仿真 Agent 测试报告",
            report_type="validation",
            instructions="记录输入、执行步骤、验证结果和待确认项。",
            project_name="仿真验证项目",
        )
        self.evidence = build_evidence_manifest(
            "测试环境完成初始化。\n执行联合仿真场景并保留日志。\n结果需要人工复核。",
            {"source_file": "case.md"},
        )
        self.spec = build_mock_spec(self.request, self.evidence, [])

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_default_render_is_valid_docx(self):
        output = self.root / "default.docx"
        metadata = render_report(self.spec, output, self.root)
        validation = validate_rendered_report(output, self.spec, self.evidence, metadata)

        self.assertEqual("default", metadata["mode"])
        self.assertTrue(validate_docx_package(output)["ok"])
        self.assertIn(validation["status"], {"passed", "passed-with-warnings"})
        self.assertIn(self.spec.title, "\n".join(p.text for p in Document(output).paragraphs))

    def test_missing_paragraph_style_fallback_does_not_duplicate_text(self):
        document = Document()
        _add_with_style(document, "single item", "Missing Style", "• ")
        self.assertEqual(["• single item"], [p.text for p in document.paragraphs])

    def test_style_reference_preserves_header_and_replaces_body(self):
        template = self.root / "reference.docx"
        doc = Document()
        doc.sections[0].header.paragraphs[0].text = "REFERENCE HEADER"
        doc.add_paragraph("REMOVE THIS SAMPLE BODY")
        doc.save(template)

        output = self.root / "styled.docx"
        metadata = render_report(self.spec, output, self.root, template)
        rendered = Document(output)
        body = "\n".join(p.text for p in rendered.paragraphs)
        header = "\n".join(p.text for p in rendered.sections[0].header.paragraphs)

        self.assertEqual("style-reference", metadata["mode"])
        self.assertIn("REFERENCE HEADER", header)
        self.assertNotIn("REMOVE THIS SAMPLE BODY", body)
        self.assertIn(self.spec.title, body)

    def test_fixed_cid629_reference_enforces_heading_body_and_toc_format(self):
        output = self.root / "strict.docx"
        metadata = render_report(
            self.spec,
            output,
            self.root,
            WORD_REPORT_TEMPLATE_PATH,
            strict_reference=True,
        )
        validation = validate_rendered_report(output, self.spec, self.evidence, metadata)
        checks = validation["strict_format"]

        self.assertEqual("strict-cid629-reference", metadata["mode"])
        self.assertEqual(14.0, checks["normal_size_pt"])
        self.assertEqual(22.0, checks["heading1_size_pt"])
        self.assertEqual(18.0, checks["heading2_size_pt"])
        self.assertGreaterEqual(checks["heading1_paragraphs"], 1)
        self.assertGreaterEqual(checks["heading2_paragraphs"], 1)
        self.assertTrue(checks["toc_levels_1_to_3"])
        self.assertTrue(checks["update_fields_on_open"])

    def test_user_format_overrides_are_applied_and_validated(self):
        output = self.root / "overridden.docx"
        options = WordFormatOptions(
            body_font="宋体",
            body_size_pt=12.5,
            line_spacing=1.25,
            first_line_indent_chars=1.5,
            heading1_font="微软雅黑",
            heading1_size_pt=20,
            heading2_font="楷体",
            heading2_size_pt=16,
            numbering_style="chinese",
        )
        metadata = render_report(
            self.spec,
            output,
            self.root,
            WORD_REPORT_TEMPLATE_PATH,
            strict_reference=True,
            format_options=options,
            format_profile_id="custom-format-test",
        )
        checks = validate_rendered_report(
            output, self.spec, self.evidence, metadata
        )["strict_format"]
        self.assertEqual("宋体", checks["normal_font"])
        self.assertEqual(12.5, checks["normal_size_pt"])
        self.assertEqual("微软雅黑", checks["heading1_font"])
        self.assertEqual(20.0, checks["heading1_size_pt"])

    def test_prepared_template_accepts_report_body_token(self):
        template = self.root / "prepared.docx"
        doc = Document()
        doc.add_heading("{{ title }}", level=1)
        doc.add_paragraph("{{p report_body }}")
        doc.save(template)

        output = self.root / "prepared-output.docx"
        metadata = render_report(self.spec, output, self.root, template)
        validation = validate_rendered_report(output, self.spec, self.evidence, metadata)

        self.assertEqual("prepared-template", metadata["mode"])
        self.assertFalse(validation["errors"])

    def test_rejects_non_zip_docx(self):
        broken = self.root / "broken.docx"
        broken.write_bytes(b"not a zip")
        with self.assertRaisesRegex(ValueError, "ZIP"):
            validate_docx_package(broken)

    def test_repairs_gbk_bytes_misread_as_latin1(self):
        original = "联合仿真测试报告"
        broken = original.encode("gb18030").decode("latin1")
        repaired, changed = normalize_mojibake(broken)
        self.assertTrue(changed)
        self.assertEqual(original, repaired)


if __name__ == "__main__":
    unittest.main()
